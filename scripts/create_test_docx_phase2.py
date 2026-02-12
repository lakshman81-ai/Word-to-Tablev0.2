from docx import Document
from docx.shared import Inches
import os

def create_complex_test_docx():
    doc = Document()
    doc.add_heading('Complex Table Scenarios (Phase 2)', 0)

    # --- Scenario 1: Partial column names ---
    doc.add_heading('1. Partial Column Names', level=1)
    doc.add_paragraph("Table with some empty header cells.")

    table1 = doc.add_table(rows=4, cols=4)
    # Header Row
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = "ID"
    hdr_cells[1].text = "" # Missing
    hdr_cells[2].text = "Value"
    hdr_cells[3].text = "" # Missing

    # Data Rows
    data = [
        ("1", "A", "100", "X"),
        ("2", "B", "200", "Y"),
        ("3", "C", "300", "Z")
    ]
    for i, row_data in enumerate(data):
        cells = table1.rows[i+1].cells
        for j, val in enumerate(row_data):
            cells[j].text = val

    # --- Scenario 2: Title in Row 1, Columns in Row 2 ---
    doc.add_heading('2. Title Row then Header Row', level=1)

    table2 = doc.add_table(rows=5, cols=3)
    # Row 1: Title (Merged?) - Simulating by just putting text in first cell or merging
    row1 = table2.rows[0]
    row1.cells[0].merge(row1.cells[2])
    row1.cells[0].text = "Sales Report 2023"

    # Row 2: Headers
    row2 = table2.rows[1]
    row2.cells[0].text = "Month"
    row2.cells[1].text = "Revenue"
    row2.cells[2].text = "Cost"

    # Data
    data2 = [("Jan", "1000", "500"), ("Feb", "1200", "600"), ("Mar", "1100", "550")]
    for i, d in enumerate(data2):
        cells = table2.rows[i+2].cells
        cells[0].text = d[0]
        cells[1].text = d[1]
        cells[2].text = d[2]

    # --- Scenario 3: Merged Rows in Data ---
    doc.add_heading('3. Merged Rows in Data', level=1)

    table3 = doc.add_table(rows=5, cols=4)
    # Header
    hdr3 = table3.rows[0].cells
    hdr3[0].text = "Item"
    hdr3[1].text = "Q1"
    hdr3[2].text = "Q2"
    hdr3[3].text = "Notes"

    # Data Row 1
    r1 = table3.rows[1].cells
    r1[0].text = "Widget A"
    r1[1].text = "10"
    r1[2].text = "20"
    r1[3].text = "Stable"

    # Data Row 2 & 3 (Merged "Notes")
    r2 = table3.rows[2].cells
    r3 = table3.rows[3].cells

    r2[0].text = "Widget B (Model 1)"
    r2[1].text = "15"
    r2[2].text = "25"

    r3[0].text = "Widget B (Model 2)"
    r3[1].text = "18"
    r3[2].text = "28"

    # Merge the last column of row 2 and 3
    r2[3].merge(r3[3])
    r2[3].text = "Merged Note for Model 1 & 2"

    # Normal Row 4
    r4 = table3.rows[4].cells
    r4[0].text = "Widget C"
    r4[1].text = "5"
    r4[2].text = "5"
    r4[3].text = "End"

    # --- Scenario 4: Staggered Headers ---
    doc.add_heading('4. Staggered Headers', level=1)
    # Col 1 name in row 2, others in row 3

    table4 = doc.add_table(rows=5, cols=3)

    # Row 1: Maybe Title or Empty
    table4.rows[0].cells[0].text = "Complex Header Table"

    # Row 2: "Category" in Col 1
    table4.rows[1].cells[0].text = "Category"

    # Row 3: "Subtype", "Count" in Col 2, 3
    table4.rows[2].cells[1].text = "Subtype"
    table4.rows[2].cells[2].text = "Count"

    # Data starts Row 4
    data4 = [("Electronics", "Phone", "50"), ("Electronics", "Laptop", "30")]
    for i, d in enumerate(data4):
        cells = table4.rows[i+3].cells
        cells[0].text = d[0]
        cells[1].text = d[1]
        cells[2].text = d[2]

    # --- Scenario 5: Gap in Header ---
    doc.add_heading('5. Gap in Header', level=1)

    table5 = doc.add_table(rows=5, cols=3)

    # Row 1: Title
    table5.rows[0].cells[0].text = "Employee List"

    # Row 2: Blank (Gap)

    # Row 3: Data starts (Implicitly? Or header?)
    # The requirement says "3 row onwards data (no column name)"
    # So Row 3 is data.

    data5 = [
        ("John Doe", "Engineer", "Active"),
        ("Jane Smith", "Manager", "Active"),
        ("Bob Jones", "Intern", "Inactive")
    ]

    for i, d in enumerate(data5):
        cells = table5.rows[i+2].cells
        cells[0].text = d[0]
        cells[1].text = d[1]
        cells[2].text = d[2]

    doc.save("test_phase2.docx")
    print("Created test_phase2.docx")

if __name__ == "__main__":
    create_complex_test_docx()
